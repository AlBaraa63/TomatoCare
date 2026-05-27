"""md_to_docx.py — convert FINAL_REPORT.md into a submission-ready Word document.

Uses markdown + BeautifulSoup + python-docx (no pandoc / docx-js required).
Produces native Word Heading 1-3 styles (so the inserted TOC field auto-populates),
bordered tables with shaded header rows, embedded figures, code blocks, lists,
blockquotes, a centred cover page, and page-numbered footers.

Run:  py ml/reports/md_to_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

HERE = Path(__file__).resolve().parent
SRC = HERE / "FINAL_REPORT.md"
OUT = HERE / "TomatoCare_AI_Report.docx"

FRONT_HEADINGS = {
    "ABSTRACT", "ACKNOWLEDGMENTS", "ACKNOWLEDGEMENTS", "TABLE OF CONTENTS",
    "LIST OF FIGURES", "LIST OF TABLES", "APPROVED BY",
    "APPROVAL FOR SUBMISSION & DECLARATION",
}


def shade(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_page_number_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for kind, text in (("begin", None), ("instr", "PAGE"), ("separate", None),
                       ("text", "1"), ("end", None)):
        if kind == "instr":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        elif kind == "text":
            el = OxmlElement("w:t")
            el.text = text
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
        run._r.append(el)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    note = OxmlElement("w:t")
    note.text = "Right-click and choose “Update Field” to build the table of contents."
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, note, end):
        run._r.append(el)


def iter_runs(node, bold=False, italic=False, code=False):
    """Yield (text, bold, italic, code) tuples from an inline HTML node tree."""
    if isinstance(node, NavigableString):
        yield (str(node), bold, italic, code)
        return
    if not isinstance(node, Tag):
        return
    if node.name == "br":
        yield ("\n", bold, italic, code)
        return
    b = bold or node.name in ("strong", "b")
    i = italic or node.name in ("em", "i")
    c = code or node.name == "code"
    for child in node.children:
        yield from iter_runs(child, b, i, c)


def fill_inline(paragraph, node) -> None:
    for text, b, i, c in iter_runs(node):
        if text == "":
            continue
        parts = text.split("\n")
        for k, part in enumerate(parts):
            if k:
                paragraph.add_run().add_break()
            run = paragraph.add_run(part)
            run.bold = b
            run.italic = i
            if c:
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)


def add_table(doc: Document, tbl: Tag) -> None:
    head = tbl.find("thead")
    headers = [th for th in head.find_all("th")] if head else []
    body_rows = [tr for tr in (tbl.find("tbody") or tbl).find_all("tr") if tr.find("td")]
    ncol = len(headers) or (len(body_rows[0].find_all("td")) if body_rows else 0)
    if ncol == 0:
        return
    table = doc.add_table(rows=0, cols=ncol)
    table.style = "Table Grid"
    table.autofit = True
    if headers:
        cells = table.add_row().cells
        for j, th in enumerate(headers):
            shade(cells[j], "D9E2F3")
            para = cells[j].paragraphs[0]
            # force bold for header
            for text, *_ in iter_runs(th):
                if text:
                    para.add_run(text).bold = True
    for tr in body_rows:
        tds = tr.find_all("td")
        cells = table.add_row().cells
        for j in range(ncol):
            para = cells[j].paragraphs[0]
            if j < len(tds):
                fill_inline(para, tds[j])
    doc.add_paragraph()  # spacer after table


def add_code_block(doc: Document, pre: Tag) -> None:
    text = pre.get_text()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    lines = text.rstrip("\n").split("\n")
    for k, line in enumerate(lines):
        if k:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_list(doc: Document, lst: Tag, level: int = 0) -> None:
    style = "List Number" if lst.name == "ol" else "List Bullet"
    if level:
        style = f"{style} {level + 1}"
    for li in lst.find_all("li", recursive=False):
        p = doc.add_paragraph(style=style)
        # inline content excluding nested lists
        for child in li.children:
            if isinstance(child, Tag) and child.name in ("ul", "ol"):
                continue
            fill_inline(p, child if isinstance(child, Tag) else child)
        for child in li.find_all(["ul", "ol"], recursive=False):
            add_list(doc, child, level + 1)


def main() -> None:
    raw = SRC.read_text(encoding="utf-8")
    raw = re.sub(r"<!--.*?-->", "", raw, flags=re.DOTALL)  # drop HTML comments
    html = markdown.markdown(raw, extensions=["tables", "fenced_code", "sane_lists"])
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()
    # base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(1))
    add_page_number_footer(sec)

    title_done = False
    seen_chapter = False
    in_cover = True
    skip_toc_stub = False

    for el in soup.children:
        if isinstance(el, NavigableString):
            continue
        if not isinstance(el, Tag):
            continue
        name = el.name
        text = el.get_text(strip=True)

        # skip the manual TOC stub between the TOC field and LIST OF FIGURES
        if skip_toc_stub:
            if name in ("h1", "h2", "h3"):
                skip_toc_stub = False
            else:
                continue

        if name == "h1":
            if not title_done:
                title_done = True
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(22)
                run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
                doc.add_paragraph()
                continue
            seen_chapter = True
            in_cover = False
            doc.add_page_break()
            doc.add_heading(text, level=1)
            continue

        if name in ("h2", "h3"):
            level = 2 if name == "h2" else 3
            if not seen_chapter:  # front matter
                up = text.upper()
                if up == "APPROVED BY":
                    in_cover = False
                if up in ("ABSTRACT", "TABLE OF CONTENTS"):
                    doc.add_page_break()
                fp = doc.add_paragraph()
                fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = fp.add_run(text)
                r.bold = True
                r.font.size = Pt(14 if level == 2 else 12)
                if up == "TABLE OF CONTENTS":
                    add_toc(doc)
                    skip_toc_stub = True
                continue
            doc.add_heading(text, level=level)
            continue

        if name == "table":
            add_table(doc, el)
            continue
        if name == "pre":
            add_code_block(doc, el)
            continue
        if name in ("ul", "ol"):
            add_list(doc, el)
            continue
        if name == "hr":
            continue
        if name == "blockquote":
            for q in el.find_all("p", recursive=False) or [el]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                for t, b, i, c in iter_runs(q):
                    if t.strip():
                        run = p.add_run(t)
                        run.italic = True
            continue

        if name == "p":
            img = el.find("img")
            if img is not None:
                src = img.get("src", "")
                path = (HERE / src).resolve()
                if path.exists():
                    doc.add_picture(str(path), width=Inches(5.8))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = (img.get("alt") or "").strip()
                if cap:
                    cp = doc.add_paragraph()
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cp.add_run(cap)
                    cr.italic = True
                    cr.font.size = Pt(9.5)
                continue
            p = doc.add_paragraph()
            if in_cover:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fill_inline(p, el)
            continue

    doc.save(str(OUT))

    # quick self-check
    d2 = Document(str(OUT))
    n_tables = len(d2.tables)
    n_paras = len(d2.paragraphs)
    n_imgs = len(d2.inline_shapes)
    print(f"[ok] wrote {OUT}")
    print(f"     paragraphs={n_paras}  tables={n_tables}  inline_images={n_imgs}")
    print(f"     size={OUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
